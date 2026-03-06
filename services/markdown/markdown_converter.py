import re
from datetime import datetime
from io import BytesIO
from config import settings

class MarkdownConverter:
    def __init__(self):
        self.default_font_name = settings.DEFAULT_FONT_NAME
        self.default_font_size = settings.DEFAULT_FONT_SIZE
    
    def convert_markdown(self, content: str, output_format: str, filename: str = None) -> BytesIO:
        """Convert Markdown to specified format and return as BytesIO"""
        if output_format == 'html':
            return self.convert_to_html(content)
        elif output_format == 'pdf':
            return self.convert_to_pdf(content)
        elif output_format == 'docx':
            return self.convert_to_docx(content)
        else:
            raise ValueError(f"Unsupported output format: {output_format}")
    
    def convert_to_html(self, content: str) -> BytesIO:
        """Convert Markdown to HTML"""
        html = self.parse_markdown(content)
        
        # Wrap in HTML document
        full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Markdown Document</title>
    <style>
        body {{
            font-family: {self.default_font_name}, Arial, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
        }}
        h1, h2, h3, h4, h5, h6 {{
            color: #333;
        }}
        code {{
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
        }}
        pre {{
            background-color: #f4f4f4;
            padding: 10px;
            border-radius: 5px;
            overflow-x: auto;
        }}
        pre code {{
            background-color: transparent;
            padding: 0;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }}
        th {{
            background-color: #f4f4f4;
        }}
    </style>
</head>
<body>
{html}
</body>
</html>"""
        
        output = BytesIO(full_html.encode('utf-8'))
        output.seek(0)
        
        return output
    
    def convert_to_pdf(self, content: str) -> BytesIO:
        """Convert Markdown to PDF (via HTML)"""
        # For now, return HTML content with PDF filename
        # In production, use a library like weasyprint or pdfkit
        html_output = self.convert_to_html(content)
        return html_output
    
    def convert_to_docx(self, content: str) -> BytesIO:
        """Convert Markdown to DOCX"""
        # Import here to avoid dependency issues
        try:
            from docx import Document
            from docx.shared import Pt
            from io import BytesIO
            
            doc = Document()
            
            # Parse and add content
            self.parse_markdown_to_docx(doc, content)
            
            # Save to BytesIO
            doc_stream = BytesIO()
            doc.save(doc_stream)
            doc_stream.seek(0)
            
            return doc_stream
        except ImportError:
            # Fallback to text format
            output = BytesIO(content.encode('utf-8'))
            output.seek(0)
            return output
    
    def parse_markdown(self, content: str) -> str:
        """Parse Markdown to HTML"""
        html = content
        
        # Code blocks
        html = re.sub(r'```(\w*)\n(.*?)```', r'<pre><code>\2</code></pre>', html, flags=re.DOTALL)
        
        # Inline code
        html = re.sub(r'`([^`]+)`', r'<code>\1</code>', html)
        
        # Headers
        html = re.sub(r'^###### (.+)$', r'<h6>\1</h6>', html, flags=re.MULTILINE)
        html = re.sub(r'^##### (.+)$', r'<h5>\1</h5>', html, flags=re.MULTILINE)
        html = re.sub(r'^#### (.+)$', r'<h4>\1</h4>', html, flags=re.MULTILINE)
        html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
        html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
        html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
        
        # Bold and italic
        html = re.sub(r'\*\*\*(.+?)\*\*\*', r'<strong><em>\1</em></strong>', html)
        html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
        html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)
        
        # Links
        html = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2">\1</a>', html)
        
        # Images
        html = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', r'<img src="\2" alt="\1">', html)
        
        # Blockquotes
        html = re.sub(r'^> (.+)$', r'<blockquote>\1</blockquote>', html, flags=re.MULTILINE)
        
        # Horizontal rules
        html = re.sub(r'^---$|^---\s$', r'<hr>', html, flags=re.MULTILINE)
        
        # Unordered lists
        html = re.sub(r'^[-*] (.+)$', r'<li>\1</li>', html, flags=re.MULTILINE)
        html = re.sub(r'((<li>.*</li>\n?)+)', r'<ul>\1</ul>', html)
        
        # Ordered lists
        html = re.sub(r'^\d+\. (.+)$', r'<li>\1</li>', html, flags=re.MULTILINE)
        
        # Tables
        html = self.convert_tables(html)
        
        # Paragraphs
        lines = html.split('\n')
        result = []
        for line in lines:
            stripped = line.strip()
            if stripped and not stripped.startswith('<'):
                result.append(f'<p>{stripped}</p>')
            else:
                result.append(line)
        html = '\n'.join(result)
        
        # Clean up empty paragraphs
        html = re.sub(r'<p>\s*</p>', '', html)
        
        return html
    
    def convert_tables(self, html: str) -> str:
        """Convert Markdown tables to HTML tables"""
        table_pattern = r'(\|.*?\n\|.*?-+\|.*?\n(?:\|.*?\n)*)'
        
        def replace_table(match):
            table_text = match.group(0)
            lines = table_text.strip().split('\n')
            
            if len(lines) < 2:
                return match.group(0)
            
            # Parse headers
            headers = [cell.strip() for cell in lines[0].split('|') if cell.strip()]
            
            # Parse data rows (skip separator)
            rows = []
            for line in lines[2:]:
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                if cells:
                    rows.append(cells)
            
            # Build HTML table
            table_html = '<table>\n<thead>\n<tr>'
            for header in headers:
                table_html += f'<th>{header}</th>'
            table_html += '</tr>\n</thead>\n<tbody>'
            
            for row in rows:
                table_html += '\n<tr>'
                for cell in row:
                    table_html += f'<td>{cell}</td>'
                table_html += '</tr>'
            
            table_html += '\n</tbody>\n</table>'
            
            return table_html
        
        return re.sub(table_pattern, replace_table, html, flags=re.MULTILINE)
    
    def parse_markdown_to_docx(self, doc, content: str):
        """Parse Markdown and add to Word document"""
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                doc.add_paragraph()
                i += 1
                continue
            
            # Headers
            if line.startswith('######'):
                doc.add_heading(line[6:].strip(), level=6)
            elif line.startswith('#####'):
                doc.add_heading(line[5:].strip(), level=5)
            elif line.startswith('####'):
                doc.add_heading(line[4:].strip(), level=4)
            elif line.startswith('###'):
                doc.add_heading(line[3:].strip(), level=3)
            elif line.startswith('##'):
                doc.add_heading(line[2:].strip(), level=2)
            elif line.startswith('#'):
                doc.add_heading(line[1:].strip(), level=1)
            
            # Bullet points
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            
            # Numbered lists
            elif re.match(r'^\d+\.\s', line):
                doc.add_paragraph(re.sub(r'^\d+\.\s', '', line), style='List Number')
            
            # Code blocks (simplified)
            elif line.startswith('```'):
                i += 1
                code_lines = []
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if code_lines:
                    para = doc.add_paragraph('\n'.join(code_lines))
                    para.runs[0].font.name = 'Courier New'
            
            # Regular paragraph
            else:
                # Process inline formatting
                text = line
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Remove bold markers
                text = re.sub(r'\*(.+?)\*', r'\1', text)  # Remove italic markers
                text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # Keep link text only
                
                doc.add_paragraph(text)
            
            i += 1
    
    def generate_filename(self, filename: str = None, output_format: str = 'html') -> str:
        """Generate a filename with timestamp if not provided"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = filename or f"markdown_{timestamp}"
        
        ext_map = {
            'html': '.html',
            'pdf': '.pdf',
            'docx': '.docx'
        }
        
        ext = ext_map.get(output_format, '.html')
        if not filename.endswith(ext):
            filename += ext
        
        return filename
