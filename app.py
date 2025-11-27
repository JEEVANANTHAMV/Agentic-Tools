from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import _Cell
import os
from datetime import datetime
import socket
import re
from typing import Optional

app = FastAPI(title="Word Document Generator API")

# Configuration
BASE_DIR = "generated_documents"
PORT = 19801

class DocumentRequest(BaseModel):
    content: str
    filename: Optional[str] = None

def get_server_ip():
    """Get server IP address"""
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except:
        return "localhost"

def create_date_folder():
    """Create folder structure based on current date"""
    today = datetime.now()
    date_path = os.path.join(
        BASE_DIR,
        today.strftime('%Y'),
        today.strftime('%m'),
        today.strftime('%d')
    )
    os.makedirs(date_path, exist_ok=True)
    return date_path

def add_page_number(doc):
    """Add page numbers to document footer"""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_element = OxmlElement(f'w:{edge}')
        edge_element.set(qn('w:val'), 'single')
        edge_element.set(qn('w:sz'), '4')
        edge_element.set(qn('w:space'), '0')
        edge_element.set(qn('w:color'), '000000')
        tcBorders.append(edge_element)
    
    tcPr.append(tcBorders)

def create_table_from_markdown(doc, table_content, default_font_name, default_font_size):
    """
    Create table from markdown syntax
    |Header1|Header2|Header3|
    |-------|-------|-------|
    |Row1Col1|Row1Col2|Row1Col3|
    |Row2Col1|Row2Col2|Row2Col3|
    """
    lines = [line.strip() for line in table_content.split('\n') if line.strip()]
    
    # Parse headers
    headers = [cell.strip() for cell in lines[0].split('|') if cell.strip()]
    
    # Parse data rows (skip separator line)
    data_rows = []
    for line in lines[2:]:  # Skip header and separator
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells:
            data_rows.append(cells)
    
    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        para = cell.paragraphs[0]
        run = para.add_run(header)
        run.bold = True
        run.font.name = default_font_name
        run.font.size = Pt(default_font_size)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        set_cell_border(cell)
    
    # Add data rows
    for row_data in data_rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            if i < len(row_cells):
                cell = row_cells[i]
                para = cell.paragraphs[0]
                
                # Process inline formatting in cell
                process_inline_formatting(para, cell_data, default_font_name, default_font_size)
                set_cell_border(cell)
    
    return table

def extract_font_settings(text):
    """Extract font settings from text if present"""
    font_pattern = r'\[FONT:([^,]+),(\d+)\]'
    font_match = re.search(font_pattern, text)
    
    if font_match:
        font_name = font_match.group(1)
        font_size = int(font_match.group(2))
        # Remove the font directive from text
        clean_text = re.sub(font_pattern, '', text).strip()
        return clean_text, font_name, font_size
    
    return text, None, None

def parse_and_format_content(doc, content):
    """
    Parse content string and apply formatting
    Supports:
    # Heading 1
    ## Heading 2
    ### Heading 3
    **bold text**
    *italic text*
    - Bullet point
    1. Numbered list
    [FONT:Arial,14]Text with custom font
    [SIZE:16]Text with custom size[/SIZE]
    |Header1|Header2| for tables
    """
    # Default font settings
    default_font_name = "Calibri"
    default_font_size = 11
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            doc.add_paragraph()
            i += 1
            continue
        
        # Extract font settings for this paragraph
        clean_line, font_name, font_size = extract_font_settings(line)
        
        # Use extracted settings or defaults
        current_font_name = font_name or default_font_name
        current_font_size = font_size or default_font_size
        
        # Check if this is a table
        if clean_line.startswith('|') and '|' in clean_line[1:]:
            # Collect all table lines
            table_lines = [line]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith('|'):
                table_lines.append(lines[j].strip())
                j += 1
            
            # Create table
            table_content = '\n'.join(table_lines)
            create_table_from_markdown(doc, table_content, current_font_name, current_font_size)
            i = j
            continue
        
        # Handle headings
        if clean_line.startswith('###'):
            heading_text = clean_line.replace('###', '').strip()
            para = doc.add_heading(heading_text, level=3)
            # Apply font to all runs in the heading
            for run in para.runs:
                run.font.name = current_font_name
                run.font.size = Pt(current_font_size)
        elif clean_line.startswith('##'):
            heading_text = clean_line.replace('##', '').strip()
            para = doc.add_heading(heading_text, level=2)
            # Apply font to all runs in the heading
            for run in para.runs:
                run.font.name = current_font_name
                run.font.size = Pt(current_font_size)
        elif clean_line.startswith('#'):
            heading_text = clean_line.replace('#', '').strip()
            para = doc.add_heading(heading_text, level=1)
            # Apply font to all runs in the heading
            for run in para.runs:
                run.font.name = current_font_name
                run.font.size = Pt(current_font_size)
        
        # Handle bullet points
        elif clean_line.startswith('- ') or clean_line.startswith('* '):
            para = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(para, clean_line[2:], current_font_name, current_font_size)
        
        # Handle numbered lists
        elif re.match(r'^\d+\.\s', clean_line):
            para = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\.\s', '', clean_line)
            process_inline_formatting(para, text, current_font_name, current_font_size)
        
        # Regular paragraph
        else:
            para = doc.add_paragraph()
            process_inline_formatting(para, clean_line, current_font_name, current_font_size)
        
        i += 1

def process_inline_formatting(para, text, default_font, default_size):
    """Process inline formatting like bold, italic, custom fonts and sizes"""
    
    # First handle [FONT:name,size]...[/FONT] blocks
    font_pattern = r'\[FONT:([^,]+),(\d+)\](.*?)\[/FONT\]'
    # Handle [SIZE:size]...[/SIZE] blocks
    size_pattern = r'\[SIZE:(\d+)\](.*?)\[/SIZE\]'
    
    current_text = text
    segments = []
    
    # Process text to identify formatted segments
    pos = 0
    while pos < len(current_text):
        # Check for FONT tag
        font_match = re.search(font_pattern, current_text[pos:])
        size_match = re.search(size_pattern, current_text[pos:])
        
        # Find the earliest match
        next_match = None
        match_type = None
        
        if font_match and (not size_match or font_match.start() <= size_match.start()):
            next_match = font_match
            match_type = 'font'
        elif size_match:
            next_match = size_match
            match_type = 'size'
        
        if next_match:
            # Add text before match
            if next_match.start() > 0:
                segments.append(('normal', current_text[pos:pos+next_match.start()], default_font, default_size))
            
            # Add matched segment
            if match_type == 'font':
                font_name = next_match.group(1)
                font_size = int(next_match.group(2))
                content = next_match.group(3)
                segments.append(('normal', content, font_name, font_size))
            else:  # size
                font_size = int(next_match.group(1))
                content = next_match.group(2)
                segments.append(('normal', content, default_font, font_size))
            
            pos += next_match.end()
        else:
            # No more matches, add remaining text
            segments.append(('normal', current_text[pos:], default_font, default_size))
            break
    
    if not segments:
        segments = [('normal', text, default_font, default_size)]
    
    # Process each segment for bold/italic
    for segment in segments:
        process_text_formatting(para, segment[1], segment[2], segment[3])

def process_text_formatting(para, text, font_name, font_size):
    """Process bold and italic formatting"""
    # Pattern for **bold**, *italic*, and ***both***
    pattern = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|[^*]+)'
    
    for match in re.finditer(pattern, text):
        segment = match.group()
        
        if not segment:
            continue
        
        run = para.add_run()
        
        # Bold and Italic ***text***
        if segment.startswith('***') and segment.endswith('***'):
            run.text = segment[3:-3]
            run.bold = True
            run.italic = True
        # Bold text **text**
        elif segment.startswith('**') and segment.endswith('**'):
            run.text = segment[2:-2]
            run.bold = True
        # Italic text *text*
        elif segment.startswith('*') and segment.endswith('*') and len(segment) > 2:
            run.text = segment[1:-1]
            run.italic = True
        else:
            run.text = segment
        
        # Apply font
        run.font.name = font_name
        run.font.size = Pt(font_size)

@app.post("/generate-document")
async def generate_document(request: DocumentRequest):
    """Generate a Word document from string content"""
    try:
        # Create date-based folder
        folder_path = create_date_folder()
        
        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = request.filename or f"document_{timestamp}"
        if not filename.endswith('.docx'):
            filename += '.docx'
        
        filepath = os.path.join(folder_path, filename)
        
        # Create document
        doc = Document()
        
        # Parse and add content
        parse_and_format_content(doc, request.content)
        
        # Add page numbers
        add_page_number(doc)
        
        # Save document
        doc.save(filepath)
        
        # Generate download URL
        server_ip = get_server_ip()
        relative_path = filepath.replace(BASE_DIR + os.sep, '').replace(os.sep, '/')
        download_url = f"http://{server_ip}:{PORT}/download/{relative_path}"
        
        return {
            "status": "success",
            "message": "Document generated successfully",
            "filename": filename,
            "filepath": filepath,
            "download_url": download_url,
            "created_at": datetime.now().isoformat()
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/download/{year}/{month}/{day}/{filename}")
async def download_file(year: str, month: str, day: str, filename: str):
    """Download generated document"""
    filepath = os.path.join(BASE_DIR, year, month, day, filename)
    
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        filepath,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        filename=filename
    )

@app.get("/")
async def root():
    """API information"""
    return {
        "name": "Word Document Generator API",
        "version": "2.0.0",
        "endpoints": {
            "generate": "/generate-document (POST)",
            "download": "/download/{year}/{month}/{day}/{filename} (GET)"
        },
        "server_ip": get_server_ip(),
        "base_dir": BASE_DIR
    }
